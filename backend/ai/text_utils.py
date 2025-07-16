import os
import re
import numpy as np
import nltk
import gensim
import gensim.models.keyedvectors as word2vec
import logging
import io

try:
    nltk.data.find("tokenizers/punkt")
except LookupError:
    nltk.download("punkt", quiet=True)

MODEL_PATH = os.path.join(os.path.dirname(__file__), "cc.ru.300.vec")

def download_fasttext_model():
    import urllib.request
    import gzip
    import shutil
    url = "https://dl.fbaipublicfiles.com/fasttext/vectors-crawl/cc.ru.300.vec.gz"
    gz_path = MODEL_PATH + ".gz"
    print("Скачиваем FastText русские вектора (1.2 ГБ)...")
    urllib.request.urlretrieve(url, gz_path)
    with gzip.open(gz_path, 'rb') as f_in:
        with open(MODEL_PATH, 'wb') as f_out:
            shutil.copyfileobj(f_in, f_out)
    os.remove(gz_path)
    print("FastText вектора скачаны и распакованы.")

def get_fasttext_model():
    if not os.path.exists(MODEL_PATH):
        download_fasttext_model()
    print("Загрузка FastText word vectors...")
    return word2vec.KeyedVectors.load_word2vec_format(MODEL_PATH, binary=False)

def clean_text(text):
    text = text.lower()
    text = re.sub(r'[^\u0430-\u044fa-z0-9\u0451\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def dynamic_fragment(text, min_words=30, max_words=120):
    from nltk.tokenize import sent_tokenize
    sentences = sent_tokenize(text, language='russian')
    chunks, buf, buf_len = [], [], 0
    for sent in sentences:
        words = sent.split()
        buf.append(sent)
        buf_len += len(words)
        if buf_len >= min_words:
            if buf_len > max_words:
                chunks.append(' '.join(buf[:-1]))
                buf = [buf[-1]]
                buf_len = len(buf[-1].split())
            else:
                chunks.append(' '.join(buf))
                buf, buf_len = [], 0
    if buf:
        chunks.append(' '.join(buf))
    return chunks

def chunk_to_vector(chunk_words, ft_model):
    if not chunk_words:
        return np.zeros(ft_model.vector_size, dtype=np.float32)
    vectors = []
    for word in chunk_words:
        if word in ft_model:
            vectors.append(ft_model[word])
        else:
            vectors.append(np.zeros(ft_model.vector_size, dtype=np.float32))
    return np.mean(vectors, axis=0)

def extract_text(file_path, pdf_timeout=180):
    import threading

    ext = os.path.splitext(file_path)[-1].lower()
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == '.docx':
        try:
            import docx
        except ImportError:
            raise ImportError("Модуль python-docx не установлен. Установите его с помощью 'pip install python-docx'")
        try:
            doc = docx.Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            raise RuntimeError(f"Ошибка чтения docx: {e}")
    elif ext == '.pdf':
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("Модуль pdfplumber не установлен. Установите его с помощью 'pip install pdfplumber'")
        result = {}
        def pdf_worker():
            try:
                with pdfplumber.open(file_path) as pdf:
                    result['text'] = "\n".join(page.extract_text() or '' for page in pdf.pages)
            except Exception as e:
                result['text'] = ''
                logging.error(f"Ошибка чтения PDF {file_path}: {e}")

        thread = threading.Thread(target=pdf_worker)
        thread.start()
        thread.join(timeout=pdf_timeout)
        if thread.is_alive():
            logging.error(f"PDF обработка {file_path} превысила таймаут {pdf_timeout} секунд. Файл пропущен.")
            return ''
        return result.get('text', '')
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def extract_text_from_bytes(content: bytes, filename: str, pdf_timeout=180):
    """
    Извлекает текст из файла, находящегося в памяти (bytes).
    """
    ext = os.path.splitext(filename)[-1].lower()
    if ext == '.txt':
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            return content.decode('cp1251', errors='ignore')
    elif ext == '.docx':
        try:
            import docx
        except ImportError:
            raise ImportError("Модуль python-docx не установлен. Установите его с помощью 'pip install python-docx'")
        try:
            file_like = io.BytesIO(content)
            doc = docx.Document(file_like)
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            logging.error(f"Ошибка чтения docx из bytes: {e}")
            return ''
    elif ext == '.pdf':
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("Модуль pdfplumber не установлен. Установите его с помощью 'pip install pdfplumber'")
        result = {}
        def pdf_worker():
            try:
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    result['text'] = "\n".join(page.extract_text() or '' for page in pdf.pages)
            except Exception as e:
                result['text'] = ''
                logging.error(f"Ошибка чтения PDF из bytes: {e}")

        import threading
        thread = threading.Thread(target=pdf_worker)
        thread.start()
        thread.join(timeout=pdf_timeout)
        if thread.is_alive():
            logging.error(f"PDF обработка (bytes) превысила таймаут {pdf_timeout} секунд. Файл пропущен.")
            return ''
        return result.get('text', '')
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def preprocess_data(file_path, ft_model, min_words=30, max_words=120):
    try:
        logging.info(f"Обработка файла: {file_path}")
        text = extract_text(file_path)
        if not text:
            logging.warning(f"Файл пустой или не удалось извлечь текст: {file_path}")
            return []
        text = clean_text(text)
        chunks = dynamic_fragment(text, min_words=min_words, max_words=max_words)
        processed_chunks = [chunk_to_vector(chunk.split(), ft_model) for chunk in chunks if chunk]
        return processed_chunks
    except Exception as e:
        logging.error(f"Ошибка при обработке файла {file_path}: {e}")
        return []